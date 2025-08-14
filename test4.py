import streamlit as st
import pandas as pd
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from io import BytesIO
import zipfile
import matplotlib.pyplot as plt
from datetime import datetime, timedelta
import unicodedata, string, re

# =======================
# Streamlit config (panggil sekali di awal)
# =======================
st.set_page_config(page_title="Surat Massal PMT", layout="centered")

# =======================
# Kamus Bahasa (pakai punyamu — tidak diubah)
# =======================
LANGUAGES = {
    "id": {
        "welcome": "Selamat Datang di Aplikasi Surat Massal PMT",
        "login": "Silakan login untuk menggunakan aplikasi ini.",
        "username": "Username",
        "password": "Password",
        "login_button": "Login",
        "logout_button": "Logout",
        "dashboard_title": "Dashboard",
        "generate_title": "Generate Surat Massal",
        "choose_language": "Pilih Bahasa / Choose Language",
        "upload_template": "Upload Template Word (.docx) — Drag & Drop atau klik",
        "upload_data": "Upload Data Excel (.xlsx) — Drag & Drop atau klik",
        "select_name_col": "Pilih kolom Nama",
        "select_link_col": "Pilih kolom Link",
        "search_name": "Cari Nama (ketik untuk filter)",
        "select_name_preview": "Pilih Nama untuk Preview",
        "hide_preview": "Sembunyikan Preview",
        "show_preview": "Tampilkan Preview",
        "preview_letter": "📖 Pratinjau Surat (Visual dan Rapi)",
        "download_preview": "⬇️ Download Preview Surat",
        "generate_all": "Generate Semua Surat",
        "processing_letters": "Sedang memproses surat...",
        "generate_done": "✅ Proses generate selesai!",
        "download_all_zip": "Download Semua Surat (ZIP)",
        "view_log": "Lihat Log Generate",
        "logout_msg": "👋 Terima Kasih!",
        "logout_submsg": "Terima kasih telah menggunakan aplikasi ini.\n\n**See you!**",
        "back_login": "🔐 Kembali ke Halaman Login",
        "total_letters": "Total Surat Dibuat",
        "letters_success": "Surat Berhasil",
        "letters_failed": "Surat Gagal",
        "templates_available": "Template Tersedia",
        "last_data_rows": "Data Peserta Terakhir",
        "letters_success_vs_failed": "Statistik Surat Berhasil vs Gagal",
        "percentage_letters": "Persentase Surat",
        "last_activity": "Aktivitas Terakhir",
        "no_data": "Belum ada data surat untuk ditampilkan.",
        "no_activity": "Belum ada aktivitas generate surat.",
        "tips": "Tips Cepat",
        "tips_content": (
            "1. Upload template dan data Excel di halaman **Generate Surat**.\n"
            "2. Pilih kolom nama dan link sesuai data.\n"
            "3. Klik **Generate Semua Surat** dan tunggu hingga selesai.\n"
            "4. Unduh file ZIP berisi surat-surat yang sudah jadi."
        ),
        "app_version": "**Versi Aplikasi:** 1.0.0",
        "no_maintenance": "⚙️ *Tidak ada pemeliharaan sistem saat ini.*",
        "upload_first": "Silakan upload template dan data Excel terlebih dahulu.",
        "login_fail": "Username atau password salah.",
    },
    "en": {
        "welcome": "Welcome to PMT Bulk Letter Application",
        "login": "Please login to use this application.",
        "username": "Username",
        "password": "Password",
        "login_button": "Login",
        "logout_button": "Logout",
        "dashboard_title": "Dashboard",
        "generate_title": "Generate Bulk Letters",
        "choose_language": "Select Language / Pilih Bahasa",
        "upload_template": "Upload Word Template (.docx) — Drag & Drop or click",
        "upload_data": "Upload Excel Data (.xlsx) — Drag & Drop or click",
        "select_name_col": "Select Name Column",
        "select_link_col": "Select Link Column",
        "search_name": "Search Name (type to filter)",
        "select_name_preview": "Select Name for Preview",
        "hide_preview": "Hide Preview",
        "show_preview": "Show Preview",
        "preview_letter": "📖 Letter Preview (Visual and Neat)",
        "download_preview": "⬇️ Download Letter Preview",
        "generate_all": "Generate All Letters",
        "processing_letters": "Processing letters...",
        "generate_done": "✅ Generation process complete!",
        "download_all_zip": "Download All Letters (ZIP)",
        "view_log": "View Generation Log",
        "logout_msg": "👋 Thank You!",
        "logout_submsg": "Thank you for using this application.\n\n**See you!**",
        "back_login": "🔐 Back to Login Page",
        "total_letters": "Total Letters Created",
        "letters_success": "Successful Letters",
        "letters_failed": "Failed Letters",
        "templates_available": "Templates Available",
        "last_data_rows": "Last Data Rows",
        "letters_success_vs_failed": "Successful vs Failed Letters Statistics",
        "percentage_letters": "Letters Percentage",
        "last_activity": "Recent Activity",
        "no_data": "No letter data to display.",
        "no_activity": "No letter generation activity yet.",
        "tips": "Quick Tips",
        "tips_content": (
            "1. Upload template and Excel data in the **Generate Letters** page.\n"
            "2. Select name and link columns as per your data.\n"
            "3. Click **Generate All Letters** and wait until completion.\n"
            "4. Download the ZIP file containing generated letters."
        ),
        "app_version": "**App Version:** 1.0.0",
        "no_maintenance": "⚙️ *No system maintenance currently.*",
        "upload_first": "Please upload the template and Excel data first.",
        "login_fail": "Wrong username or password.",
    },
}

def t(key):
    return LANGUAGES.get(st.session_state.lang, LANGUAGES["id"]).get(key, key)

# =======================
# Preview DOCX (punyamu, tidak diubah)
# =======================
def render_docx_preview_visual(doc):
    st.subheader(t("preview_letter"))
    style = """
    <style>
        .docx-preview {
            background: #fff;
            padding: 20px;
            border: 1px solid #ddd;
            border-radius: 8px;
            font-family: Arial, sans-serif;
            font-size: 15px;
            line-height: 1.6;
            text-align: justify;
            max-height: 400px;
            overflow-y: auto;
        }
        .docx-preview p { margin-bottom: 1em; }
        .docx-preview a { color: #1a0dab; text-decoration: underline; }
        .docx-preview strong { font-weight: bold; }
        .docx-preview em { font-style: italic; }
    </style>
    """
    html = '<div class="docx-preview">'
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        run_html = ""
        for run in p.runs:
            text = run.text.replace("\n", "<br>")
            if run.bold and run.italic:
                run_html += f"<strong><em>{text}</em></strong>"
            elif run.bold:
                run_html += f"<strong>{text}</strong>"
            elif run.italic:
                run_html += f"<em>{text}</em>"
            else:
                run_html += text
        html += f"<p>{run_html}</p>"
    html += "</div>"
    st.markdown(style + html, unsafe_allow_html=True)

# =======================
# Helpers
# =======================
def safe_filename_base(name: str):
    valid = f"-_.() {string.ascii_letters}{string.digits}"
    norm = unicodedata.normalize("NFKD", str(name)).encode("ascii", "ignore").decode()
    cleaned = "".join(c for c in norm if c in valid).strip()
    return cleaned or "surat"

def _clear_paragraph(p):
    for r in p.runs:
        r._r.getparent().remove(r._r)

def replace_placeholder_with_hyperlink(p, text_before, url, text_after):
    _clear_paragraph(p)
    if text_before:
        rb = p.add_run(text_before); rb.font.name = "Arial"; rb.font.size = Pt(12)

    if url:
        r_id = p.part.relate_to(url, RT.HYPERLINK, is_external=True)
        hyperlink = OxmlElement("w:hyperlink"); hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")

        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Arial"); rFonts.set(qn("w:hAnsi"), "Arial")
        rPr.append(rFonts)

        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "24"); rPr.append(sz)
        color = OxmlElement("w:color"); color.set(qn("w:val"), "0000FF"); rPr.append(color)
        underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single"); rPr.append(underline)

        new_run.append(rPr)
        text_elem = OxmlElement("w:t"); text_elem.set(qn("xml:space"), "preserve"); text_elem.text = str(url)
        new_run.append(text_elem)
        hyperlink.append(new_run)
        p._p.append(hyperlink)
    else:
        rp = p.add_run("(tautan tidak tersedia)"); rp.font.name = "Arial"; rp.font.size = Pt(12)

    if text_after:
        ra = p.add_run(text_after); ra.font.name = "Arial"; ra.font.size = Pt(12)

    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def normalize_url(u: str):
    u = ("" if pd.isna(u) else str(u)).strip()
    if not u:
        return ""
    if not (u.startswith("http://") or u.startswith("https://")):
        u = "https://" + u
    return u

# =======================
# Generate (pakai kolom No + prefix file)
# =======================
def generate_letters_with_progress(template_file, df, col_name, col_link):
    df[col_link] = df[col_link].map(normalize_url)

    output_zip = BytesIO()
    log = []
    pad_width = st.session_state.get("pad_width", len(str(len(df))))
    used_names = set()

    with zipfile.ZipFile(output_zip, "w") as zf:
        total = len(df)
        progress_bar = st.progress(0)
        status_text = st.empty()

        for i, (idx, row) in enumerate(df.iterrows(), start=1):
            try:
                tpl = DocxTemplate(template_file)
                tpl.render({
                    "nama_penyelenggara": row[col_name],
                    "short_link": "[short_link]",
                    "no": int(row["No"]),
                })
                temp_buf = BytesIO(); tpl.save(temp_buf); temp_buf.seek(0)

                doc = Document(temp_buf)
                pattern = re.escape("[short_link]")
                for p in doc.paragraphs:
                    if re.search(pattern, p.text):
                        parts = p.text.split("[short_link]")
                        before = parts[0] if len(parts) > 0 else ""
                        after  = parts[1] if len(parts) > 1 else ""
                        replace_placeholder_with_hyperlink(p, before, row[col_link], after)

                # styling umum (skip paragraf yg punya hyperlink)
                for p in doc.paragraphs:
                    if p._p.xpath(".//w:hyperlink"):  # biar hyperlink tetap biru/underline
                        continue
                    for run in p.runs:
                        run.font.name = "Arial"; run.font.size = Pt(12)

                final_buf = BytesIO(); doc.save(final_buf)

                base = safe_filename_base(row[col_name])
                prefix = f"{int(row['No']):0{pad_width}d}"
                candidate = f"{prefix} - {base}.docx"
                n = 1
                while candidate.lower() in used_names:
                    n += 1
                    candidate = f"{prefix} - {base} ({n}).docx"
                used_names.add(candidate.lower())

                zf.writestr(candidate, final_buf.getvalue())
                log.append({"Nama": row[col_name], "Status": "✅ Berhasil"})
            except Exception as e:
                log.append({"Nama": row.get(col_name, '(unknown)'), "Status": f"❌ Gagal: {e}"})

            progress = int(i / total * 100)
            progress_bar.progress(progress)
            status_text.text(f"{t('processing_letters')} {i} / {total}")

    output_zip.seek(0)
    st.session_state.generate_log = st.session_state.get("generate_log", []) + log
    st.session_state.template_count = 1
    st.session_state.last_data_rows = len(df)
    return output_zip, log

# =======================
# Session & Pages
# =======================
SESSION_TIMEOUT = timedelta(minutes=15)

def check_session_timeout():
    if "last_active" in st.session_state:
        if datetime.now() - st.session_state.last_active > SESSION_TIMEOUT:
            st.session_state.clear()
            st.rerun()
    st.session_state.last_active = datetime.now()

def page_generate():
    st.title(t("generate_title"))
    template_file = st.file_uploader(t("upload_template"), type="docx")
    data_file = st.file_uploader(t("upload_data"), type="xlsx")

    if template_file and data_file:
        try:
            df = pd.read_excel(data_file)

            # Tambah kolom No otomatis di paling depan
            if "No" not in df.columns:
                df.insert(0, "No", range(1, len(df) + 1))
            st.session_state.pad_width = len(str(len(df)))

            st.success(f"{len(df)} rows loaded successfully")
            st.dataframe(df)

            col_name = st.selectbox(t("select_name_col"), df.columns)
            col_link = st.selectbox(t("select_link_col"), df.columns)

            search_name = st.text_input(t("search_name"), "")
            filtered_names = df[df[col_name].astype(str).str.contains(search_name, case=False, na=False)][col_name].unique()
            selected_name = st.selectbox(t("select_name_preview"), filtered_names)

            st.session_state.df = df
            st.session_state.col_name = col_name
            st.session_state.col_link = col_link
            st.session_state.selected_name = selected_name
            st.session_state.template_file = template_file

        except Exception as e:
            st.error(f"Failed to read Excel file: {e}")
            return

        if 'show_preview' not in st.session_state:
            st.session_state.show_preview = True
        if st.button(t("hide_preview") if st.session_state.show_preview else t("show_preview")):
            st.session_state.show_preview = not st.session_state.show_preview

        if st.session_state.show_preview and selected_name:
            row = df[df[col_name] == selected_name].iloc[0]
            tpl = DocxTemplate(template_file)
            tpl.render({
                "nama_penyelenggara": row[col_name],
                "short_link": "[short_link]",
                "no": int(row["No"]),
            })
            temp_buf = BytesIO(); tpl.save(temp_buf); temp_buf.seek(0)

            doc = Document(temp_buf)
            for p in doc.paragraphs:
                if "[short_link]" in p.text:
                    parts = p.text.split("[short_link]")
                    before = parts[0] if len(parts) > 0 else ""
                    after  = parts[1] if len(parts) > 1 else ""
                    replace_placeholder_with_hyperlink(p, before, row[col_link], after)

            render_docx_preview_visual(doc)

            preview_buf = BytesIO(); doc.save(preview_buf); preview_buf.seek(0)
            st.download_button(
                label=f"{t('download_preview')} ({row[col_name]})",
                data=preview_buf.getvalue(),
                file_name=f"preview_{row[col_name]}.docx",
            )

        if st.button(t("generate_all")):
            with st.spinner(t("processing_letters")):
                zip_file, log = generate_letters_with_progress(template_file, df, col_name, col_link)
            st.success(t("generate_done"))
            st.toast(t("generate_done"), icon="✅")
            st.download_button(t("download_all_zip"), zip_file.getvalue(), file_name="surat_massal.zip")
            with st.expander(t("view_log")):
                st.dataframe(pd.DataFrame(log))
    else:
        st.info(t("upload_first"))

def page_home():
    st.title(t("dashboard_title"))
    st.markdown(f"{t('welcome')}, **{st.session_state.username}**!")

    with st.expander(t("tips"), expanded=True):
        st.write(t("tips_content"))

    st.markdown("---")

    generate_log = st.session_state.get("generate_log", [])
    total_surat = len(generate_log)
    berhasil = sum(1 for item in generate_log if item["Status"].startswith("✅"))
    gagal = total_surat - berhasil

    template_tersedia = st.session_state.get("template_count", 1)
    data_peserta_terakhir = st.session_state.get("last_data_rows", 0)

    statistik_data = {
        "Statistik": [
            t("total_letters"),
            t("letters_success"),
            t("letters_failed"),
            t("templates_available"),
            t("last_data_rows"),
        ],
        "Jumlah": [
            total_surat, berhasil, gagal, template_tersedia, data_peserta_terakhir,
        ],
    }
    df_statistik = pd.DataFrame(statistik_data)
    st.markdown("### " + t("dashboard_title"))
    st.table(df_statistik)

    st.markdown("---")
    st.markdown("### " + t("letters_success_vs_failed"))
    fig, ax = plt.subplots()
    ax.bar([t("letters_success"), t("letters_failed")], [berhasil, gagal], color=["green", "red"])
    ax.set_ylabel(t("total_letters"))
    ax.set_title(t("letters_success_vs_failed"))
    st.pyplot(fig); plt.close(fig)

    st.markdown("---")
    st.markdown("### " + t("percentage_letters"))
    fig2, ax2 = plt.subplots()
    if total_surat > 0:
        ax2.pie(
            [berhasil, gagal],
            labels=[t("letters_success"), t("letters_failed")],
            autopct="%1.1f%%",
            colors=["green", "red"],
            startangle=90,
            wedgeprops={"edgecolor": "black"},
        )
        ax2.axis("equal")
        st.pyplot(fig2); plt.close(fig2)
    else:
        st.write(t("no_data"))

    st.markdown("---")
    st.markdown("### " + t("last_activity"))
    aktivitas = [{"Aktivitas": f"{t('generate_title')} untuk {item['Nama']}", "Status": item["Status"]}
                 for item in reversed(generate_log[-5:])]
    if aktivitas:
        st.table(pd.DataFrame(aktivitas))
    else:
        st.write(t("no_activity"))

    st.markdown("---")
    st.markdown(t("app_version"))
    st.markdown(t("no_maintenance"))

# =======================
# Login & Shell
# =======================
def show_login():
    st.title(t("welcome"))
    st.markdown(t("login"))
    with st.form("login_form"):
        username = st.text_input(t("username"))
        password = st.text_input(t("password"), type="password")
        submitted = st.form_submit_button(t("login_button"))
        if submitted:
            if username == "admin" and password == "surat123":
                st.session_state.login_state = True
                st.session_state.username = username
                st.session_state.logout_message = False
                st.rerun()
            else:
                st.error(t("login_fail"))

def show_main_app():
    check_session_timeout()
    st.sidebar.success(f"{t('welcome')}, {st.session_state.username}")
    if st.sidebar.button(t("logout_button")):
        st.session_state.logout_message = True
        st.session_state.login_state = False
        st.session_state.username = ""
        st.rerun()

    st.sidebar.title(t("choose_language"))
    lang = st.sidebar.selectbox("", ["id", "en"],
                                index=0 if st.session_state.get("lang", "id") == "id" else 1,
                                format_func=lambda x: "Indonesia" if x == "id" else "English")
    st.session_state.lang = lang

    st.sidebar.title("Menu")
    page = st.sidebar.radio("Navigasi", [t("dashboard_title"), t("generate_title")])
    if page == t("dashboard_title"):
        page_home()
    else:
        page_generate()

if "login_state" not in st.session_state:
    st.session_state.login_state = False
if "lang" not in st.session_state:
    st.session_state.lang = "id"

if st.session_state.get("logout_message", False):
    st.title(t("logout_msg"))
    st.markdown(t("logout_submsg"))
    if st.button(t("back_login")):
        st.session_state.logout_message = False
        st.rerun()
elif st.session_state.login_state:
    show_main_app()
else:
    show_login()
