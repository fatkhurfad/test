import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import zipfile

st.set_page_config(page_title="Generator Surat Massal", layout="centered")
st.title("üìÑ Generator Surat Massal")
# üñåÔ∏è Load Modern CSS
with open("style.css") as f:
    st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    rPr.append(sz)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    new_run.append(rPr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

uploaded_template = st.file_uploader("Upload Template Word (.docx)", type="docx")
uploaded_excel = st.file_uploader("Upload Data Excel (.xlsx)", type="xlsx")

if uploaded_template and uploaded_excel:
    df = pd.read_excel(uploaded_excel)

    if len(df.columns) >= 2:
        st.subheader("üìå Pilih kolom untuk placeholder:")
        col_nama = st.selectbox("Ganti {{nama_penyelenggara}} dengan:", df.columns)
        col_link = st.selectbox("Ganti {{short_link}} dengan:", df.columns)

        file_name_format = st.text_input(
            "üìù Format nama file output (gunakan {{nama_penyelenggara}})",
            value="Surat - {{nama_penyelenggara}}"
        )

        # Validasi isi template
        required_placeholders = ["{{nama_penyelenggara}}", "{{short_link}}"]
        doc_check = Document(uploaded_template)
        doc_text = "\n".join([p.text for p in doc_check.paragraphs])
        missing_placeholders = [ph for ph in required_placeholders if ph not in doc_text]
        if missing_placeholders:
            st.warning(f"‚ö†Ô∏è Template tidak mengandung placeholder: {', '.join(missing_placeholders)}")

        if col_nama not in df.columns or col_link not in df.columns:
            st.error("‚ùå Kolom Excel tidak valid.")
            st.stop()

        row_index = st.number_input("üîç Pilih baris untuk preview (mulai dari 1)", min_value=1, max_value=len(df), value=1)

        if st.button("Tampilkan Preview"):
            row = df.iloc[row_index - 1]
            doc = Document(uploaded_template)

            for p in doc.paragraphs:
                for run in p.runs:
                    if "{{nama_penyelenggara}}" in run.text:
                        run.text = run.text.replace("{{nama_penyelenggara}}", str(row[col_nama]))

            for p in doc.paragraphs:
                if "{{short_link}}" in p.text:
                    parts = p.text.split("{{short_link}}")
                    p.clear()
                    if parts[0]: p.add_run(parts[0])
                    add_hyperlink(p, str(row[col_link]), str(row[col_link]))
                    if len(parts) > 1: p.add_run(parts[1])

            for p in doc.paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(12)

            preview_text = "\n".join([p.text for p in doc.paragraphs])
            st.text_area("üìù Isi Surat Preview:", value=preview_text, height=400)

            preview_buffer = BytesIO()
            doc.save(preview_buffer)
            preview_buffer.seek(0)
            st.download_button(
                label="üìÑ Download Preview Surat",
                data=preview_buffer.getvalue(),
                file_name=f"preview_{row[col_nama]}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        if st.button("üîÑ Generate Semua Surat"):
            output_zip = BytesIO()
            failed = []
            success = 0

            with zipfile.ZipFile(output_zip, "w") as zf:
                for idx, row in df.iterrows():
                    try:
                        doc = Document(uploaded_template)

                        for p in doc.paragraphs:
                            for run in p.runs:
                                if "{{nama_penyelenggara}}" in run.text:
                                    run.text = run.text.replace("{{nama_penyelenggara}}", str(row[col_nama]))

                        for p in doc.paragraphs:
                            if "{{short_link}}" in p.text:
                                parts = p.text.split("{{short_link}}")
                                p.clear()
                                if parts[0]: p.add_run(parts[0])
                                add_hyperlink(p, str(row[col_link]), str(row[col_link]))
                                if len(parts) > 1: p.add_run(parts[1])

                        for p in doc.paragraphs:
                            for run in p.runs:
                                run.font.name = "Arial"
                                run.font.size = Pt(12)

                        custom_filename = file_name_format.replace("{{nama_penyelenggara}}", str(row[col_nama]))
                        filename = f"{custom_filename.replace('/', '-')}.docx"

                        buffer = BytesIO()
                        doc.save(buffer)
                        zf.writestr(filename, buffer.getvalue())
                        success += 1
                    except Exception as e:
                        failed.append((idx + 1, str(row[col_nama]), str(e)))

            st.success(f"‚úÖ Surat berhasil dibuat: {success}")
            if failed:
                st.error(f"‚ùå Gagal dibuat: {len(failed)}")
                for item in failed:
                    st.text(f"Baris {item[0]} ({item[1]}): {item[2]}")

            st.download_button(
                label="üì• Download ZIP Semua Surat",
                data=output_zip.getvalue(),
                file_name="surat_massal_output.zip",
                mime="application/zip"
            )
    else:
        st.warning("Excel harus memiliki setidaknya 2 kolom.")
